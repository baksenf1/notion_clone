from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "ASSIGNMENT_4_REPORT.docx"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, True)
        shade_cell(hdr[i], "D9EAF7")
        if widths:
            hdr[i].width = widths[i]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = widths[i]
    document.add_paragraph()
    return table


def add_heading(document, text, level=1):
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(31, 78, 121)
    return paragraph


def add_body(document, text):
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.line_spacing = 1.08
    return paragraph


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.7)
section.right_margin = Inches(0.7)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(10)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Software Development Case Studies\nAssignment 4")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(31, 78, 121)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Product Thinking & MVP Development\nCollaborative Workspace MVP")
run.font.size = Pt(12)
run.italic = True

add_table(
    doc,
    ["Field", "Value"],
    [
        ["Course", "Software Development Case Study"],
        ["Week", "Week 8-9"],
        ["Topic", "Product Thinking, Ethics & Product Debt"],
        ["Student(s)", "[Your name / team names]"],
        ["Prototype", "Local repository / GitHub link: [insert link]"],
    ],
    [Inches(1.5), Inches(5.7)],
)

add_heading(doc, "Part 1 - MVP Definition", 1)
add_heading(doc, "1.1 MVP Hypothesis", 2)
add_table(
    doc,
    ["Field", "Answer"],
    [
        ["The problem you are solving", "Small student project teams lose time because project notes, task ownership, deadlines, and progress are scattered across messengers, documents, spreadsheets, and calendar tools."],
        ["Your riskiest assumption", "The riskiest assumption is that small teams will find value in one lightweight shared workspace enough to create and update tasks there, instead of continuing to coordinate informally in Telegram/WhatsApp and Google Docs. If teams do not return to update tasks after the first setup, the product has no real value."],
        ["Your MVP hypothesis", "If we build a collaborative workspace where a student team can create a project, assign tasks with status, priority, and due date, and view recent progress in one dashboard, we believe 2-5 person student project teams will use it at least twice during a project week to coordinate work, because seeing ownership and deadlines in one shared place reduces coordination friction."],
        ["What the MVP is NOT", "The MVP is not a full Notion clone. It does not include advanced document blocks, real-time multiplayer editing, file storage, AI writing, complex automations, payments, mobile apps, or deep calendar integrations. These are left out because they do not test the riskiest assumption: whether teams will actually coordinate around a shared workspace and task flow."],
        ["How you will test it", "The minimum test is an end-to-end workflow: user signs in, creates/opens a workspace, creates a project, creates a task with assignee/status/priority/due date, and sees the task appear in the task table and dashboard. The test is successful if a target user can complete this flow without help and understands who owns what work by the end."],
    ],
    [Inches(1.8), Inches(5.5)],
)
add_body(doc, "Eric Ries defines an MVP as the smallest product version that enables a full Build-Measure-Learn loop with minimum effort. This prototype follows that idea because it does not try to ship a smaller version of every future feature. It focuses only on the workflow that tests the core value assumption: shared coordination.")

add_heading(doc, "1.2 Product-Market Fit Signal", 2)
add_table(
    doc,
    ["Question", "Answer"],
    [
        ["Who is your primary user?", "The primary user is a university student working in a 2-5 person software/project team for coursework, hackathons, or semester assignments. This user needs lightweight coordination but usually does not want the complexity of enterprise project management tools."],
        ["What retention signal will you track?", "The main retention signal is weekly task maintenance: a workspace has at least two active members who create or update tasks on two or more separate days in the same project week. This is stronger than signup because it shows the team returned to the product when coordination was actually needed."],
        ["Distribution risk", "The strongest distribution risk comes from Notion, Trello, Jira, Google Workspace, Microsoft Teams, Moodle, and Telegram/WhatsApp communities. Any of them could add or already has task/project features. Similar to the Kaspi bill-splitting example from the lecture, a platform with existing daily user traffic can add a small feature and instantly reach users. This would hurt the product if our only value were basic task creation. It would not completely kill the MVP if we focus on a student-specific workflow: course project workspace, simple member roles, deadline visibility, and low setup friction for small teams."],
    ],
    [Inches(1.8), Inches(5.5)],
)

add_heading(doc, "1.3 Product Debt Awareness", 2)
add_table(
    doc,
    ["Debt Risk", "How It Could Appear in Your Product", "How You Will Avoid It"],
    [
        ["Feature creep", "The team could start adding document editing, chat, file uploads, AI summaries, calendar sync, custom dashboards, and mobile support before proving that users care about the core workspace-task flow.", "Keep the MVP scope limited to workspace, project, members, pages, and tasks. New features are postponed unless they directly help test whether users coordinate work in the product."],
        ["Wrong early-adopter assumption", "The first users may be software students who already understand task boards and are more tolerant of unfinished tools. They may not represent less technical student teams who need simpler language and fewer configuration choices.", "Test with at least two types of student teams: software/project management students and non-technical coursework teams. Ask whether they would actually return during a real assignment."],
        ["North star metric misalignment", "A tempting metric would be number of created tasks, pages, or time spent in the app. This could encourage busywork rather than useful coordination.", "Use weekly active task updates by at least two members and percentage of tasks moved to Done before due date as better signals. These measure coordination and completion, not just activity."],
    ],
    [Inches(1.45), Inches(3.0), Inches(2.85)],
)

add_heading(doc, "Part 2 - Ethics & Privacy", 1)
add_heading(doc, "2.1 Privacy by Design", 2)
add_table(
    doc,
    ["Principle", "How Your Product Applies It"],
    [
        ["Data minimisation - collect only what you use", "The MVP only needs account identity, workspace membership, project names, page content, task details, assignment, status, priority, and due date. It does not need location, phone number, demographic data, private calendar access, or message history. Optional profile images are only used for member recognition in the UI."],
        ["Purpose limitation - data collected for one feature cannot be reused for another", "Task and workspace data is used only for collaboration features: displaying projects, assigning work, filtering tasks, and showing progress. It will not be reused for ranking students, advertising, external profiling, or academic evaluation without explicit consent."],
        ["Default = private - users opt in to sharing, not out", "Workspaces are private by default. A user must be invited through a workspace invite flow before they can access workspace projects, pages, members, and tasks. The product should not make projects public by default and should not expose workspace data through unauthenticated routes."],
        ["Right to erasure (GDPR Art. 17) - full deletion including logs and backups", "A user should be able to request deletion of their account and personal data. In the MVP, workspace deletion removes workspace projects, pages, members, and tasks. For a production version, deletion must also cover application logs, analytics records, and backup retention windows. Logs should avoid storing page/task content so erasure is realistic."],
    ],
    [Inches(2.25), Inches(5.05)],
)
add_body(doc, "Privacy by Design also connects to GDPR Article 25: privacy should be built into the architecture rather than added later. In this prototype, the architectural decision is to scope data by workspace and require authentication before workspace data can be accessed.")

add_heading(doc, "2.2 Ethical Risk Analysis", 2)
add_table(
    doc,
    ["Risk", "EAD Principle", "Who Is Affected", "Your Mitigation"],
    [
        ["Unequal visibility could turn the product into a hidden surveillance tool. If the app later adds activity tracking, it could pressure students or allow team leaders to judge classmates by raw activity instead of contribution quality.", "Human Well-being; Transparency; Accountability", "Students, especially quieter team members or students with limited availability.", "Do not add hidden productivity scoring. If activity indicators are added, make them visible to all affected users and explain what is measured. Measure task outcomes and ownership clarity rather than time online. Apply the Greyball test: if a tracking feature must be hidden from students to work, it should not be built."],
        ["Invite codes and member roles could expose private coursework or allow unauthorized access if links are shared outside the team.", "Awareness of Misuse; Competence", "Team members, course groups, and anyone whose project notes or tasks are stored in the workspace.", "Keep workspaces private by default, validate membership on backend routes, allow invite code reset, and use role-based permissions for sensitive actions. The product should make it clear who has access to a workspace and should not rely only on frontend hiding."],
    ],
    [Inches(2.25), Inches(1.5), Inches(1.55), Inches(2.0)],
)

add_heading(doc, "Part 3 - MVP Prototype", 1)
add_heading(doc, "3.1 Prototype Requirements", 2)
add_table(
    doc,
    ["Component", "Requirement", "How the Prototype Meets It"],
    [
        ["Core Feature", "At least one end-to-end workflow that directly tests the MVP hypothesis. Input -> Processing -> Output. No fake buttons.", "The implemented workflow is: create/open workspace -> create project -> create task -> assign member, status, priority, and due date -> backend validates and saves task -> frontend updates task list/dashboard."],
        ["User Interface", "Minimum 2 screens with functional navigation. UI must reflect your target user.", "The prototype includes authentication screens, workspace dashboard, members page, projects/project details, pages, tasks table, and settings. Navigation is provided through the app layout/sidebar."],
        ["Data Flow", "Data must change based on user input. Mock data is acceptable.", "User input is sent from React forms to the Express API, stored in MongoDB, and returned to the frontend through API queries. Tasks can be filtered by project, status, priority, assignee, keyword, and due date."],
    ],
    [Inches(1.3), Inches(2.7), Inches(3.3)],
)
add_body(doc, "The prototype is intentionally focused on depth over breadth. The main tested feature is not many productivity tools; it is the full coordination flow from creating work to making it visible to the team.")

add_heading(doc, "3.2 Hypothesis Test Report", 2)
add_table(
    doc,
    ["Question", "Answer"],
    [
        ["Did the prototype test your MVP hypothesis?", "Yes. The hypothesis was: If we build a collaborative workspace where a student team can create a project, assign tasks with status, priority, and due date, and view recent progress in one dashboard, we believe 2-5 person student project teams will use it at least twice during a project week to coordinate work, because seeing ownership and deadlines in one shared place reduces coordination friction. The prototype tests the first half of this hypothesis by allowing users to complete the core workflow end-to-end. It does not yet prove long-term retention because that requires real team usage over time."],
        ["What did you learn that you did not know before building?", "Building the prototype showed that workspace collaboration is too broad as a product promise. The real value is clearer when the product is framed around task ownership and deadlines. It also showed that permissions and invite flows are not secondary details; they are part of the core trust model for collaboration."],
        ["What product debt did you introduce, and why?", "Some product debt was introduced by including pages, workspace analytics, members, projects, and tasks in the same MVP. This risks making the product look like a general Notion clone. The reason for accepting this debt was that workspace context, project context, and task assignment are connected in the coordination workflow. However, advanced page editing and automation should remain outside the MVP until task coordination is validated."],
        ["What is the next hypothesis to test?", "The next hypothesis is: If we add lightweight reminders and a clearer weekly progress summary, we believe student teams will update tasks more consistently before deadlines, because the current dashboard alone may not create enough return motivation. This should be tested only after measuring whether teams already return to update task status without reminders."],
    ],
    [Inches(1.9), Inches(5.4)],
)

add_heading(doc, "Part 4 - Reading Connection", 1)
add_table(
    doc,
    ["Source (title, author, year)", "Key idea you applied", "Where it appears in your work"],
    [
        ["The Lean Startup - Eric Ries, 2011", "An MVP should enable a Build-Measure-Learn loop with the least effort, rather than being a simple incomplete version of the final product.", "This idea shaped the scope of the prototype. Instead of building a full Notion clone with chat, file storage, AI, and real-time editing, the MVP focuses on one measurable learning goal: whether student teams will coordinate projects through shared workspace tasks. The PMF signal also follows this logic by tracking repeated task updates, not signups or downloads."],
    ],
    [Inches(1.85), Inches(2.35), Inches(3.1)],
)

add_heading(doc, "Conclusion", 1)
add_body(doc, "This MVP tests whether small student teams need a lightweight shared workspace for project coordination. The prototype provides the minimum workflow needed to test that value: authentication, private workspace, project creation, task creation, assignment, status, priority, due date, dashboard visibility, and task filtering. The strongest business risk is distribution because existing tools such as Notion, Trello, Google Workspace, Microsoft Teams, Moodle, and Telegram could absorb basic task coordination. Therefore, the product should avoid competing as a generic task tool and instead focus on the specific needs of student project teams.")
add_body(doc, "The main ethical responsibility is to keep collaboration transparent and private. The product should help teams coordinate work, not secretly monitor students or expose private coursework. Following Privacy by Design and IEEE EAD principles helps keep those decisions architectural from the beginning.")

doc.save(OUT)
print(OUT)
